import os
from pathlib import Path

import pytest
from docx import Document as DocxDocument

from app.llm_parser import _load_doc_text


def test__load_doc_text_extracts_and_prints_text(tmp_path: Path):
    # Create a temporary .docx with paragraphs and a table
    doc = DocxDocument()
    doc.add_paragraph("First paragraph")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "C"
    table.cell(1, 1).text = "D"

    tmp_file = tmp_path / "sample.docx"
    doc.save(tmp_file)

    # Run extraction
    extracted = _load_doc_text(str(tmp_file))

    # Output the converted text so it is visible when running with -s
    print(extracted)

    # Basic assertions about content and structure
    assert "First paragraph" in extracted
    assert "A | B" in extracted
    assert "C | D" in extracted
    assert extracted.strip() == extracted


