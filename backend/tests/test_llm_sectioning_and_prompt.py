from pathlib import Path
from unittest.mock import patch

from docx import Document as DocxDocument

from app.llm_parser import extract_rows_from_docx


def _make_doc_with_sections(tmp_path: Path) -> str:
    doc = DocxDocument()
    # Country header boundary
    doc.add_paragraph("(Spain)")
    doc.add_paragraph("Project Alpha status updates...")
    doc.add_paragraph("")
    # Uppercase heading boundary
    doc.add_paragraph("WEEKLY HIGHLIGHTS")
    doc.add_paragraph("Project Beta progress is on track.")
    # Blank line block boundary
    doc.add_paragraph("")
    doc.add_paragraph("Misc notes for Project Gamma.")

    path = tmp_path / "sample_sections.docx"
    doc.save(path)
    return str(path)


def test_extract_rows_calls_llm_per_section_and_merges(tmp_path: Path):
    file_path = _make_doc_with_sections(tmp_path)

    # Prepare 3 synthetic responses (one per section), each with one row
    def _fake_chat(payload_messages, **kwargs):  # noqa: ARG001
        # Always return a single entry JSON
        return {
            "choices": [
                {
                    "message": {
                        "content": "{\n  \"rows\": [{\n    \"project_name\": \"Alpha\", \n    \"title\": \"Alpha - CW01\", \n    \"summary\": \"Some summary.\", \n    \"next_actions\": null, \n    \"owner\": null, \n    \"category\": \"Development\", \n    \"source_text\": \"Original sentence.\"\n  }]\n}"
                    }
                }
            ]
        }

    with patch("app.llm_parser._azure_chat_completion", side_effect=_fake_chat) as mocked:
        rows = extract_rows_from_docx(file_path, "CW01", "DEV")

    # We expect at least 3 calls (country line, uppercase heading, blank block)
    assert mocked.call_count >= 3
    assert isinstance(rows, list)
    assert len(rows) >= 3
    for r in rows:
        assert r.get("summary")
        assert r.get("source_text")
        assert len(r.get("summary")) <= 1000


