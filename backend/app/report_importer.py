"""
Report importer functions for processing DOCX files and saving to database.
"""
import hashlib
import logging
import os
import re
from datetime import datetime, date, timedelta
from pathlib import Path
from typing import Dict, List, Optional, Callable

from docx import Document
from sqlalchemy import text
from sqlalchemy.orm import Session

from .llm_parser import extract_rows_from_docx
from .utils import parse_filename, get_project_code_by_name_db, seed_projects_from_csv, parse_docx_rows

logger = logging.getLogger(__name__)


def _calculate_file_sha256(file_path: str) -> str:
    """Calculate SHA256 hash of a file."""
    sha256_hash = hashlib.sha256()
    with open(file_path, "rb") as f:
        for byte_block in iter(lambda: f.read(4096), b""):
            sha256_hash.update(byte_block)
    return sha256_hash.hexdigest()


def _get_cw_wednesday_date(year: int, cw: int) -> date:
    """Get the Wednesday date for a given calendar week.
    
    This ensures that projects are stored in the target year for easier querying.
    For example, 2025 CW01 projects will be stored on 2025-01-01 (Wednesday)
    instead of 2024-12-30 (Monday).
    
    Args:
        year: Year (e.g., 2025)
        cw: Calendar week number (1-53)
    
    Returns:
        date: The Wednesday of that calendar week
    """
    # January 4th is always in week 1
    jan4 = date(year, 1, 4)
    # Find the Monday of week 1
    week1_monday = jan4 - timedelta(days=jan4.weekday())
    # Calculate the Wednesday of the target week (Monday + 2 days)
    target_wednesday = week1_monday + timedelta(weeks=cw - 1, days=2)
    return target_wednesday


def _create_or_get_upload_record(
    db: Session,
    file_path: str,
    original_filename: str,
    created_by: str,
    force_import: bool = False
) -> Dict:
    """Create or get existing upload record based on SHA256 hash."""
    # Calculate file hash
    file_hash = _calculate_file_sha256(file_path)
    file_size = os.path.getsize(file_path)
    
    # Parse filename to extract metadata
    try:
        year, cw_label, raw_category, category = parse_filename(original_filename)
        cw_num = int(cw_label[2:])  # Extract number from "CW01" -> 1
        log_date = _get_cw_wednesday_date(year, cw_num)
    except ValueError:
        # Fallback for files that don't match expected pattern
        cw_label = None
        category = "Unknown"
        log_date = date.today()
    
    # Check if upload already exists by SHA256 (always check to avoid unique errors)
    existing = db.execute(
        text("SELECT id, status FROM report_uploads WHERE sha256 = :sha256"),
        {"sha256": file_hash}
    ).first()

    if existing:
        logger.info(f"File already uploaded with ID: {existing.id}")
        # If force_import, reuse existing upload record to allow reprocessing
        # If not forcing, also return existing to signal duplicate
        return {
            "upload_id": existing.id,
            "is_new": False,
            "cw_label": cw_label,
            "category": category,
            "log_date": log_date
        }
    
    # Create new upload record
    result = db.execute(
        text("""
            INSERT INTO report_uploads (
                original_filename, storage_path, mime_type, file_size_bytes,
                sha256, status, cw_label, created_by, updated_by
            ) VALUES (
                :original_filename, :storage_path, :mime_type, :file_size_bytes,
                :sha256, :status, :cw_label, :created_by, :updated_by
            )
            RETURNING id
        """),
        {
            "original_filename": original_filename,
            "storage_path": file_path,
            "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "file_size_bytes": file_size,
            "sha256": file_hash,
            "status": "received",
            "cw_label": cw_label,
            "created_by": created_by,
            "updated_by": created_by,
        }
    )
    
    upload_id = result.scalar()
    logger.info(f"Created new upload record with ID: {upload_id}")
    
    return {
        "upload_id": upload_id,
        "is_new": True,
        "cw_label": cw_label,
        "category": category,
        "log_date": log_date
    }


def import_single_docx_simple_with_metadata(
    db: Session,
    file_path: str,
    original_filename: str,
    created_by: str,
    force_import: bool = False
) -> Dict:
    """
    Import a single DOCX file using simple parsing (no LLM) and save to database.
    
    Args:
        db: Database session
        file_path: Path to the DOCX file
        original_filename: Original filename for metadata extraction
        created_by: User who initiated the import
        
    Returns:
        Dict with upload_id and rows_created count
    """
    logger.info(f"Starting simple import of {original_filename}")
    
    # Ensure projects are seeded from CSV
    try:
        seed_projects_from_csv(db, created_by=created_by)
    except Exception as e:
        logger.warning(f"Could not seed projects: {e}")
    
    # Create or get upload record
    upload_info = _create_or_get_upload_record(db, file_path, original_filename, created_by, force_import)
    upload_id = upload_info["upload_id"]
    cw_label = upload_info["cw_label"]
    category = upload_info["category"]
    log_date = upload_info["log_date"]
    
    # If this is an existing upload, check if we should skip processing
    if not upload_info["is_new"] and not force_import:
        # Count existing project history records
        existing_count = db.execute(
            text("SELECT COUNT(*) FROM project_history WHERE source_upload_id = :upload_id"),
            {"upload_id": upload_id}
        ).scalar()
        
        if existing_count > 0:
            logger.info(f"Upload {upload_id} already has {existing_count} project history records")
            return {
                "upload_id": upload_id,
                "rows_created": 0  # No new rows created
            }
    
    # Parse document using simple text extraction
    rows_created = 0
    
    try:
        # Use the more sophisticated parsing logic from utils.py
        # Create a mock UploadFile object to use the existing parser
        class MockFile:
            def __init__(self, file_path):
                self.file = open(file_path, 'rb')
        
        mock_file = MockFile(file_path)
        
        try:
            # Parse using the sophisticated logic that handles project identification
            parsed_rows = parse_docx_rows(mock_file, cw_label or "CW01", category)
            
            # Convert to project sections format
            project_sections = []
            for row in parsed_rows:
                project_name = row.get("title", "Unknown Project")
                if project_name and " - " in project_name:
                    project_name = project_name.split(" - ")[0]  # Remove " - CW01" suffix
                source_text = row.get("summary", "")
                if source_text.strip():
                    project_sections.append((project_name, source_text))
        finally:
            mock_file.file.close()
        
        if not project_sections:
            # Fallback to simple parsing if sophisticated parsing fails
            document = Document(file_path)
            full_text = []
            
            # Extract paragraphs
            for paragraph in document.paragraphs:
                para_text = (paragraph.text or "").strip()
                if para_text:
                    full_text.append(para_text)
            
            # Extract table content
            for table in document.tables:
                for row in table.rows:
                    cells = [(cell.text or "").strip() for cell in row.cells]
                    cells = [c for c in cells if c]
                    if cells:
                        full_text.append(" | ".join(cells))
            
            if not full_text:
                logger.warning("No text content found in document")
                combined_text = ""
            else:
                combined_text = "\n".join(full_text)
            
            # Try to identify project sections using simple heuristics
            project_sections = _extract_project_sections_simple(combined_text)
            
            if not project_sections:
                # Create a single entry with all content
                project_sections = [("Unknown Project", combined_text)]
        
        # Create project history records
        for project_name, source_text in project_sections:
            # Try to find project code
            project_code = get_project_code_by_name_db(db, project_name)
            
            if not project_code:
                # Create virtual project code
                project_code = f"VIRT_{project_name.upper().replace(' ', '_')[:20]}"
                # Ensure it's unique in projects table
                counter = 1
                base_code = project_code
                while True:
                    existing = db.execute(
                        text("SELECT 1 FROM projects WHERE project_code = :code LIMIT 1"),
                        {"code": project_code}
                    ).first()
                    if not existing:
                        break
                    counter += 1
                    project_code = f"{base_code}_{counter}"
                
                # Create the virtual project in projects table
                db.execute(
                    text("""
                        INSERT INTO projects (
                            project_code, project_name, status, created_by, updated_by
                        ) VALUES (
                            :project_code, :project_name, :status, :created_by, :updated_by
                        )
                    """),
                    {
                        "project_code": project_code,
                        "project_name": project_name,
                        "status": 1,  # Active status
                        "created_by": created_by,
                        "updated_by": created_by,
                    }
                )
            
            # Check for duplicates (same project_code + log_date + source_upload_id)
            # Note: DB has unique constraint on (project_code, log_date, category)
            existing_record = db.execute(
                text("""
                    SELECT id FROM project_history 
                    WHERE project_code = :project_code 
                    AND log_date = :log_date 
                    AND source_upload_id = :source_upload_id
                """),
                {
                    "project_code": project_code,
                    "log_date": log_date,
                    "source_upload_id": upload_id
                }
            ).first()
            
            if existing_record:
                logger.info(f"Skipping duplicate record for {project_code} on {log_date}")
                continue
            
            # Insert project history record
            db.execute(
                text("""
                    INSERT INTO project_history (
                        project_code, category, entry_type, log_date, cw_label,
                        title, summary, source_text, source_upload_id,
                        created_by, updated_by
                    ) VALUES (
                        :project_code, :category, :entry_type, :log_date, :cw_label,
                        :title, :summary, :source_text, :source_upload_id,
                        :created_by, :updated_by
                    )
                """),
                {
                    "project_code": project_code,
                    "category": category,
                    "entry_type": "Report",
                    "log_date": log_date,
                    "cw_label": cw_label,
                    "title": f"{project_name} - {cw_label}" if cw_label else project_name,
                    "summary": source_text[:5000],  # Limit summary length
                    "source_text": source_text,
                    "source_upload_id": upload_id,
                    "created_by": created_by,
                    "updated_by": created_by,
                }
            )
            rows_created += 1
        
        # Update upload status
        db.execute(
            text("""
                UPDATE report_uploads 
                SET status = 'parsed', parsed_at = NOW(), updated_by = :updated_by
                WHERE id = :upload_id
            """),
            {"upload_id": upload_id, "updated_by": created_by}
        )
        
        db.commit()
        logger.info(f"Successfully imported {rows_created} project records")
        
        return {
            "upload_id": upload_id,
            "rows_created": rows_created
        }
        
    except Exception as e:
        logger.error(f"Error importing document: {e}")
        db.rollback()
        
        # Update upload status to failed
        try:
            db.execute(
                text("""
                    UPDATE report_uploads 
                    SET status = 'failed', updated_by = :updated_by
                    WHERE id = :upload_id
                """),
                {"upload_id": upload_id, "updated_by": created_by}
            )
            db.commit()
        except Exception:
            pass
        
        raise


def _extract_project_sections_simple(text: str) -> List[tuple[str, str]]:
    """Extract project sections using simple pattern matching."""
    sections = []
    
    # Common project patterns observed in documents
    project_patterns = [
        r"(?:^|\n)\s*[•·]\s*([^:\n]+):\s*([^\n•]*(?:\n(?!\s*[•·])[^\n•]*)*)",
        r"(?:^|\n)\s*\(([^)]+)\)\s*([A-Za-z][^:\n]*?)(?:\s+\d+(?:\.\d+)?\s*MW)?\s*:?\s*([^\n]*(?:\n(?!\s*[•·(])[^\n]*)*)",
        r"(?:^|\n)\s*([A-Za-z][A-Za-z0-9\s\-_]+?)\s*:?\s*([^\n]*(?:\n(?!\s*[•·(A-Z])[^\n]*)*)",
    ]
    
    for pattern in project_patterns:
        matches = re.finditer(pattern, text, re.MULTILINE | re.IGNORECASE)
        for match in matches:
            if len(match.groups()) >= 2:
                project_name = match.group(1).strip()
                content = match.group(2).strip() if len(match.groups()) > 1 else ""
                
                # Add remaining content if available
                if len(match.groups()) > 2 and match.group(3):
                    content += " " + match.group(3).strip()
                
                # Skip very short project names or content
                if len(project_name) > 2 and len(content) > 10:
                    sections.append((project_name, content))
    
    # If no sections found, try splitting by bullet points
    if not sections:
        lines = text.split('\n')
        current_project = None
        current_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Check if this looks like a project header
            if (line.startswith('•') or line.startswith('·') or 
                re.match(r'^[A-Za-z][A-Za-z0-9\s\-_]+:', line) or
                re.match(r'^\([^)]+\)\s*[A-Za-z]', line)):
                
                # Save previous project if exists
                if current_project and current_content:
                    sections.append((current_project, '\n'.join(current_content)))
                
                # Start new project
                current_project = re.sub(r'^[•·]\s*', '', line)
                current_project = re.sub(r':\s*$', '', current_project)
                current_content = []
            else:
                # Add to current project content
                if current_project:
                    current_content.append(line)
        
        # Don't forget the last project
        if current_project and current_content:
            sections.append((current_project, '\n'.join(current_content)))
    
    return sections


def import_single_docx_llm_with_metadata(
    db: Session,
    file_path: str,
    original_filename: str,
    created_by: str,
    project_code_mapper: Optional[Callable[[str], Optional[str]]] = None,
    force_import: bool = False
) -> Dict:
    """
    Import a single DOCX file using LLM parsing and save to database.
    
    Args:
        db: Database session
        file_path: Path to the DOCX file
        original_filename: Original filename for metadata extraction
        created_by: User who initiated the import
        project_code_mapper: Optional function to map project names to codes
        
    Returns:
        Dict with upload_id and rows_created count
    """
    logger.info(f"Starting LLM import of {original_filename}")
    
    # Ensure projects are seeded from CSV
    try:
        seed_projects_from_csv(db, created_by=created_by)
    except Exception as e:
        logger.warning(f"Could not seed projects: {e}")
    
    # Create or get upload record
    upload_info = _create_or_get_upload_record(db, file_path, original_filename, created_by, force_import)
    upload_id = upload_info["upload_id"]
    cw_label = upload_info["cw_label"]
    category = upload_info["category"]
    log_date = upload_info["log_date"]
    
    # If this is an existing upload, check if we should skip processing
    if not upload_info["is_new"] and not force_import:
        existing_count = db.execute(
            text("SELECT COUNT(*) FROM project_history WHERE source_upload_id = :upload_id"),
            {"upload_id": upload_id}
        ).scalar()
        
        if existing_count > 0:
            logger.info(f"Upload {upload_id} already has {existing_count} project history records")
            return {
                "upload_id": upload_id,
                "rows_created": 0
            }
    
    rows_created = 0
    
    try:
        # Extract data using LLM
        raw_category = category.upper() if category != "Unknown" else "DEV"
        llm_rows = extract_rows_from_docx(file_path, cw_label or "CW01", raw_category)
        
        if not llm_rows:
            logger.warning("No data extracted by LLM")
            llm_rows = []
        
        # Process extracted rows
        for row_data in llm_rows:
            project_name = row_data.get("project_name", "Unknown Project")
            
            # Map project name to code
            if project_code_mapper:
                project_code = project_code_mapper(project_name)
            else:
                project_code = get_project_code_by_name_db(db, project_name)
            
            if not project_code:
                # Create auto-generated project code
                project_code = f"AUTO_{project_name.upper().replace(' ', '_')[:20]}"
                counter = 1
                base_code = project_code
                while True:
                    existing = db.execute(
                        text("SELECT 1 FROM projects WHERE project_code = :code LIMIT 1"),
                        {"code": project_code}
                    ).first()
                    if not existing:
                        break
                    counter += 1
                    project_code = f"{base_code}_{counter}"
                
                # Create the auto project in projects table
                db.execute(
                    text("""
                        INSERT INTO projects (
                            project_code, project_name, status, created_by, updated_by
                        ) VALUES (
                            :project_code, :project_name, :status, :created_by, :updated_by
                        )
                    """),
                    {
                        "project_code": project_code,
                        "project_name": project_name,
                        "status": 1,  # Active status
                        "created_by": created_by,
                        "updated_by": created_by,
                    }
                )
            
            # Check for duplicates (same project_code + log_date + source_upload_id)
            # Note: DB has unique constraint on (project_code, log_date, category)
            existing_record = db.execute(
                text("""
                    SELECT id FROM project_history 
                    WHERE project_code = :project_code 
                    AND log_date = :log_date 
                    AND source_upload_id = :source_upload_id
                """),
                {
                    "project_code": project_code,
                    "log_date": log_date,
                    "source_upload_id": upload_id
                }
            ).first()
            
            if existing_record:
                logger.info(f"Skipping duplicate record for {project_code} on {log_date}")
                continue
            
            # Insert project history record
            summary = row_data.get("summary", "")
            source_text = row_data.get("source_text", summary)
            
            db.execute(
                text("""
                    INSERT INTO project_history (
                        project_code, category, entry_type, log_date, cw_label,
                        title, summary, source_text, next_actions, owner,
                        source_upload_id, created_by, updated_by
                    ) VALUES (
                        :project_code, :category, :entry_type, :log_date, :cw_label,
                        :title, :summary, :source_text, :next_actions, :owner,
                        :source_upload_id, :created_by, :updated_by
                    )
                """),
                {
                    "project_code": project_code,
                    "category": row_data.get("category", category),
                    "entry_type": row_data.get("entry_type", "Report"),
                    "log_date": log_date,
                    "cw_label": row_data.get("cw_label", cw_label),
                    "title": row_data.get("title", f"{project_name} - {cw_label}"),
                    "summary": summary[:5000],  # Limit summary length
                    "source_text": source_text,
                    "next_actions": row_data.get("next_actions"),
                    "owner": row_data.get("owner"),
                    "source_upload_id": upload_id,
                    "created_by": created_by,
                    "updated_by": created_by,
                }
            )
            rows_created += 1
        
        # Update upload status
        db.execute(
            text("""
                UPDATE report_uploads 
                SET status = 'parsed', parsed_at = NOW(), updated_by = :updated_by
                WHERE id = :upload_id
            """),
            {"upload_id": upload_id, "updated_by": created_by}
        )
        
        db.commit()
        logger.info(f"Successfully imported {rows_created} project records using LLM")
        
        return {
            "upload_id": upload_id,
            "rows_created": rows_created
        }
        
    except Exception as e:
        logger.error(f"Error importing document with LLM: {e}")
        db.rollback()
        
        # Update upload status to failed
        try:
            db.execute(
                text("""
                    UPDATE report_uploads 
                    SET status = 'failed', updated_by = :updated_by
                    WHERE id = :upload_id
                """),
                {"upload_id": upload_id, "updated_by": created_by}
            )
            db.commit()
        except Exception:
            pass
        
        raise


def import_single_docx(
    db: Session,
    file_path: str,
    default_project_code: str,
    created_by: str
) -> Dict:
    """
    Simple import function for backward compatibility.
    
    Args:
        db: Database session
        file_path: Path to the DOCX file
        default_project_code: Default project code to use for all entries
        created_by: User who initiated the import
        
    Returns:
        Dict with upload_id and rows_created count
    """
    filename = Path(file_path).name
    return import_single_docx_simple_with_metadata(db, file_path, filename, created_by)


def import_single_docx_llm(
    db: Session,
    file_path: str,
    project_code_mapper: Optional[Callable[[str], Optional[str]]] = None,
    created_by: str = "system"
) -> Dict:
    """
    LLM import function for backward compatibility.
    
    Args:
        db: Database session
        file_path: Path to the DOCX file
        project_code_mapper: Optional function to map project names to codes
        created_by: User who initiated the import
        
    Returns:
        Dict with upload_id and rows_created count
    """
    filename = Path(file_path).name
    return import_single_docx_llm_with_metadata(db, file_path, filename, created_by, project_code_mapper)
