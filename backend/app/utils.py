"""
Utility functions shared across modules
"""
import csv
import logging
import re
from pathlib import Path
from typing import Optional

from docx import Document
from fastapi import UploadFile
from sqlalchemy import text

logger = logging.getLogger(__name__)

# Legacy strict pattern (kept for reference)
_FILENAME_RE_STRICT = re.compile(r"^(?P<year>\d{4})_CW(?P<cw>\d{2})_(?P<cat>DEV|EPC|FINANCE|INVESTMENT)\.docx$", re.IGNORECASE)

# Flexible patterns for CW and category extraction
_CW_PATTERN = re.compile(r"CW(\d{1,2})", re.IGNORECASE)
_CATEGORY_PATTERNS = [
    (re.compile(r"(?:^|[^a-zA-Z])(DEV|DEVELOPMENT)(?:[^a-zA-Z]|$)", re.IGNORECASE), "DEV"),
    (re.compile(r"(?:^|[^a-zA-Z])(EPC)(?:[^a-zA-Z]|$)", re.IGNORECASE), "EPC"),
    (re.compile(r"(?:^|[^a-zA-Z])(FINANCE|FINANCIAL|FIN)(?:[^a-zA-Z]|$)", re.IGNORECASE), "FINANCE"),
    (re.compile(r"(?:^|[^a-zA-Z])(INVESTMENT|INVEST|INV)(?:[^a-zA-Z]|$)", re.IGNORECASE), "INVESTMENT"),
]

_CAT_MAP = {
    "DEV": "Development",
    "EPC": "EPC",
    "FINANCE": "Finance",
    "INVESTMENT": "Investment",
}


def parse_filename(filename: str):
    """Parse filename to extract year, cw_label and category.
    
    This function is now more flexible and can handle various filename formats
    as long as they contain:
    1. A calendar week number (CW##)
    2. A category (DEV, EPC, FINANCE, INVESTMENT or their variations)
    
    Examples of supported formats:
    - 2025_CW01_DEV.docx (strict format)
    - Weekly Report_CW16 - DEV.docx
    - CW01 Development Report.docx
    - Report CW02 EPC 2025.docx
    """
    # Remove .docx extension for parsing
    name_without_ext = filename.lower()
    if name_without_ext.endswith('.docx'):
        name_without_ext = name_without_ext[:-5]
    
    # Try strict pattern first (for backward compatibility)
    strict_match = _FILENAME_RE_STRICT.match(filename)
    if strict_match:
        year = int(strict_match.group("year"))
        cw_label = f"CW{strict_match.group('cw').upper()}"
        raw = strict_match.group("cat").upper()
        category = _CAT_MAP.get(raw, raw)
        return year, cw_label, raw, category
    
    # Extract CW number
    cw_match = _CW_PATTERN.search(filename)
    if not cw_match:
        raise ValueError("INVALID_NAME: No calendar week (CW##) found in filename")
    
    cw_num = int(cw_match.group(1))
    cw_label = f"CW{cw_num:02d}"  # Format as CW01, CW02, etc.
    
    # Extract category
    category_raw = None
    category = None
    for pattern, cat_raw in _CATEGORY_PATTERNS:
        if pattern.search(filename):
            category_raw = cat_raw
            category = _CAT_MAP.get(cat_raw, cat_raw)
            break
    
    if not category_raw:
        raise ValueError("INVALID_NAME: No valid category (DEV, EPC, FINANCE, INVESTMENT) found in filename")
    
    # Try to extract year (optional, default to current year if not found)
    year_match = re.search(r"\b(20\d{2})\b", filename)
    if year_match:
        year = int(year_match.group(1))
    else:
        # Default to current year if no year found in filename
        from datetime import datetime
        year = datetime.now().year
    
    return year, cw_label, category_raw, category


def parse_docx_rows(file: UploadFile, cw_label: str, category: str) -> list[dict]:
    """Enhanced project-aware parser that identifies project sections and aggregates text by project.
    
    Identifies project start lines using multiple patterns:
    1. • project_name: (bullet with colon)
    2. • project_name (bullet on separate line)
    3. (Country/Region) project_name capacity/number (e.g., "(Spain) Taurus A-3 105MW")
    4. Standalone project_name/number line (e.g., "Taurus A-1")
    
    Aggregates source_text from project start line until next project or end of file.
    """
    rows: list[dict] = []
    try:
        document = Document(file.file)
    except Exception:
        # parsing failed -> return one empty row to avoid crashing
        rows.append({
            "category": category,
            "entry_type": "Report",
            "cw_label": cw_label,
            "title": None,
            "summary": "",
            "next_actions": None,
            "owner": None,
            "attachment_url": None,
        })
        return rows

    # Load project names for matching
    project_names = _load_all_project_names()
    
    # Extract all lines in order
    all_lines: list[str] = []
    for p in document.paragraphs:
        text = (p.text or "").strip()
        if text:
            all_lines.append(text)
    
    # Also extract table content as lines
    for table in document.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                all_lines.append(" | ".join(cells))

    if not all_lines:
        # No content found, return empty summary
        rows.append({
            "category": category,
            "entry_type": "Report", 
            "cw_label": cw_label,
            "title": None,
            "summary": "",
            "next_actions": None,
            "owner": None,
            "attachment_url": None,
        })
        return rows

    # Process lines to identify project sections
    projects = _identify_project_sections(all_lines, project_names)
    
    # Convert projects to rows
    for project_name, source_text in projects:
        if source_text.strip():  # Only create rows with content
            rows.append({
                "category": category,
                "entry_type": "Report",
                "cw_label": cw_label,
                "title": f"{project_name} - {cw_label}",
                "summary": source_text.strip(),
                "next_actions": None,
                "owner": None,
                "attachment_url": None,
            })
    
    # If no projects found, create one summary row
    if not rows:
        summary = "\n".join(all_lines).strip()
        rows.append({
            "category": category,
            "entry_type": "Report",
            "cw_label": cw_label,
            "title": None,
            "summary": summary,
            "next_actions": None,
            "owner": None,
            "attachment_url": None,
        })
    
    return rows


def _load_all_project_names() -> set[str]:
    """Load all project names from CSV for project identification (case-insensitive)."""
    try:
        mapping = load_project_name_to_code_mapping()
        # Get original project names (before lowercase normalization)
        csv_path = _default_project_csv_path()
        project_names = set()

        if csv_path.exists():
            # Try different encodings to handle special characters
            for encoding in ["utf-8", "latin-1", "cp1252"]:
                try:
                    with csv_path.open("r", encoding=encoding) as f:
                        reader = csv.DictReader(f, delimiter=";")
                        for row in reader:
                            name = (row.get("project_name") or "").strip()
                            status_str = (row.get("status") or "0").strip()
                            if name and status_str in ("1", "true", "True"):
                                project_names.add(name.lower())  # Convert to lowercase
                    break  # Successfully loaded with this encoding
                except UnicodeDecodeError:
                    continue  # Try next encoding

        return project_names
    except Exception as e:
        logger.warning(f"Failed to load project names: {e}")
        return set()


def _identify_project_sections(all_lines: list[str], project_names: set[str]) -> list[tuple[str, str]]:
    """Identify project sections by searching for known project names, project patterns, and special formats.

    Returns list of (project_name, aggregated_source_text) tuples.
    """
    # Join all lines into a single text for easier searching
    full_text = "\n".join(all_lines)
    text_lower = full_text.lower()

    # Find all project name occurrences with their positions
    project_positions = []

    # 1. Search for exact project names from CSV
    for project_name in project_names:
        name_lower = project_name.lower()

        # Find all occurrences of this project name
        start = 0
        while True:
            pos = text_lower.find(name_lower, start)
            if pos == -1:
                break

            # Validate that this is a word boundary match (not part of another word)
            if _is_valid_project_match(full_text, pos, len(project_name)):
                project_positions.append((pos, project_name))

            start = pos + 1

    # 2. Search for common project patterns in the document
    project_patterns = _find_project_patterns_in_text(full_text, project_names)

    for pos, pattern_name in project_patterns:
        project_positions.append((pos, pattern_name))

    # 3. Search for special format patterns like "(Country) ProjectName capacity"
    special_formats = _find_special_format_projects(full_text)

    for pos, project_name, country, capacity in special_formats:
        project_positions.append((pos, project_name))

    # Sort by position in the document
    project_positions.sort(key=lambda x: x[0])

    # Remove duplicates (keep first occurrence of each project, prefer exact matches)
    seen_projects = set()
    unique_positions = []
    for pos, name in project_positions:
        # Create a normalized key for deduplication (handle similar projects)
        # Remove capacity information and extra spaces for better deduplication
        normalized_name = re.sub(r'\s+\d+(?:\.\d+)?\s*(?:MW|MVA|MWp|MWh|MWh/2h|MWh/4h|MWh/6h|MWh/8h|MWh/10h|MWh/12h|MWh/400MWh|MWh/100MW)+?$', '', name.lower().strip())
        normalized_name = re.sub(r'\s+', ' ', normalized_name).strip()

        if normalized_name not in seen_projects:
            unique_positions.append((pos, name))
            seen_projects.add(normalized_name)

    if not unique_positions:
        return []

    # Extract text segments between project positions
    projects = []
    for i, (pos, project_name) in enumerate(unique_positions):
        # Find the start position for this project's text
        start_pos = pos

        # Find the end position (start of next project or end of text)
        if i + 1 < len(unique_positions):
            end_pos = unique_positions[i + 1][0]
        else:
            end_pos = len(full_text)

        # Extract the text segment for this project
        project_text = full_text[start_pos:end_pos].strip()

        if project_text:
            projects.append((project_name, project_text))

    return projects


def _is_valid_project_match(text: str, pos: int, length: int) -> bool:
    """Validate that a project name match is not part of another word."""
    # Convert text to lowercase for case-insensitive matching
    text_lower = text.lower()

    # Check character before the match
    if pos > 0:
        char_before = text_lower[pos - 1]
        if char_before.isalnum():
            return False

    # Check character after the match
    end_pos = pos + length
    if end_pos < len(text_lower):
        char_after = text_lower[end_pos]
        if char_after.isalnum():
            return False

    return True


def _find_special_format_projects(full_text: str) -> list[tuple[int, str, str, str]]:
    """Find special format projects like '(Country) ProjectName capacity'.

    Returns: list of (position, project_name, country, capacity)
    """
    special_projects = []

    # Multiple patterns to handle various formats
    patterns = [
        # Standard format: (Country) ProjectName capacity
        r'\(([^)]+)\)\s*([A-Za-z_0-9\s\-+]+?)\s+(\d+(?:\.\d+)?\s*(?:MW|MVA|MWp|MWh|MWh/2h|MWh/4h|MWh/6h|MWh/8h|MWh/10h|MWh/12h|MWh/400MWh|MWh/100MW)+?)',

        # Format without spaces before capacity: (Country) ProjectName capacity
        r'\(([^)]+)\)\s*([A-Za-z_0-9\s\-+]+?)(\d+(?:\.\d+)?\s*(?:MW|MVA|MWp|MWh|MWh/2h|MWh/4h|MWh/6h|MWh/8h|MWh/10h|MWh/12h|MWh/400MWh|MWh/100MW)+?)',

        # Format with complex capacity: (Country) ProjectName capacity/unit
        r'\(([^)]+)\)\s*([A-Za-z_0-9\s\-+]+?)\s+(\d+(?:\.\d+)?(?:MW|MVA|MWp|MWh|MWh/2h|MWh/4h|MWh/6h|MWh/8h|MWh/10h|MWh/12h)+?/\d+(?:\.\d+)?(?:MW|MVA|MWp|MWh|MWh/2h|MWh/4h|MWh/6h|MWh/8h|MWh/10h|MWh/12h|MWh/400MWh|MWh/100MW)+?)',

        # Simple format: (Country) ProjectName (no capacity)
        r'\(([^)]+)\)\s*([A-Za-z_0-9\s\-+]+?)(?:\s*\([^)]*\))?\s*$',
    ]

    for pattern in patterns:
        for match in re.finditer(pattern, full_text, re.IGNORECASE | re.MULTILINE):
            pos = match.start()
            country = match.group(1).strip()
            project_part = match.group(2).strip()
            capacity = match.group(3).strip() if len(match.groups()) > 2 and match.group(3) else ""

            # Clean up project name (remove extra spaces, normalize)
            project_name = re.sub(r'\s+', ' ', project_part).strip()

            # Skip if project name is too short or contains only numbers/special chars
            if len(project_name) < 2 or re.match(r'^[\d\s\-\+]+$', project_name):
                continue

            special_projects.append((pos, project_name, country, capacity))

    return special_projects


def _find_project_patterns_in_text(full_text: str, project_names: set[str]) -> list[tuple[int, str]]:
    """Find project patterns in text that might not exactly match CSV names (case-insensitive)."""
    patterns = []
    text_lower = full_text.lower()

    # Common project patterns observed in the document (already lowercase)
    project_patterns = [
        "taurus a-1", "taurus a-2", "taurus a-3",
        "peñaflor", "mudejar", "andujar", "almodovar",
        "hermann", "carmona", "brovales", "cabrovales",
        "tordesillas", "zaratan_bess", "divor"
    ]

    for pattern in project_patterns:
        start = 0
        while True:
            pos = text_lower.find(pattern, start)
            if pos == -1:
                break

            # Validate word boundary
            if _is_valid_project_match(full_text, pos, len(pattern)):
                # Use the original case from the document
                original_text = full_text[pos:pos + len(pattern)]
                patterns.append((pos, original_text))

            start = pos + 1

    # Look for patterns like "(Country) ProjectName XX.XMW" - case insensitive
    import re
    country_pattern = re.compile(r'\(([^)]+)\)\s*([A-Z][A-Za-z_0-9\s\-]+?)\s*(?:\d+(?:\.\d+)?\s*MW)', re.IGNORECASE)

    for match in country_pattern.finditer(full_text):
        pos = match.start()
        country = match.group(1).strip()
        project_part = match.group(2).strip()

        # Clean up the project name
        project_name = project_part.strip()
        if project_name and len(project_name) > 3:
            patterns.append((pos, project_name))

    # Look for standalone project names followed by capacity - case insensitive
    standalone_pattern = re.compile(r'^([A-Z][A-Za-z_0-9\s\-]+?)\s+(\d+(?:\.\d+)?\s*MW):', re.MULTILINE | re.IGNORECASE)

    for match in standalone_pattern.finditer(full_text):
        pos = match.start()
        project_name = match.group(1).strip()

        if project_name and len(project_name) > 3:
            patterns.append((pos, project_name))

    return patterns


def _detect_project_start_line(line: str, project_names: set[str]) -> str | None:
    """Detect if a line indicates the start of a project section.
    
    Returns the detected project name or None.
    """
    line_stripped = line.strip()
    if not line_stripped:
        return None
    
    # Pattern 1: • project_name: (bullet with colon)
    bullet_colon_match = re.match(r'^\s*[•·‧▪▫]\s*([^:]+):\s*$', line_stripped)
    if bullet_colon_match:
        candidate = bullet_colon_match.group(1).strip()
        if _is_known_project(candidate, project_names):
            return candidate
    
    # Pattern 2: • project_name (bullet on separate line)
    bullet_match = re.match(r'^\s*[•·‧▪▫]\s*(.+)$', line_stripped)
    if bullet_match:
        candidate = bullet_match.group(1).strip()
        if _is_known_project(candidate, project_names):
            return candidate
    
    # Pattern 3: (Country/Region) project_name capacity/number
    # e.g., "(Spain) Taurus A-3 105MW", "(Spain) Hermann 103.9 MW"
    country_match = re.match(r'^\s*\([^)]+\)\s*(.+?)(?:\s+\d+(?:\.\d+)?\s*MW|\s+\d+MW|\s*\d+(?:\.\d+)?\s*MW)?$', line_stripped)
    if country_match:
        candidate = country_match.group(1).strip()
        if _is_known_project(candidate, project_names):
            return candidate
        # Also try variations like "Taurus A-3" from "Taurus A-3 105MW"
        name_parts = candidate.split()
        if len(name_parts) >= 2:
            for i in range(2, len(name_parts) + 1):
                partial = " ".join(name_parts[:i])
                if _is_known_project(partial, project_names):
                    return partial
    
    # Pattern 4: Standalone project_name/number line
    # e.g., "Taurus A-1", "Peñaflor 105.1MW"
    # Look for lines with project-like patterns (spaces, hyphens, numbers)
    if re.search(r'[A-Za-z]+[\s\-_]+[A-Za-z0-9]+', line_stripped):
        # Remove common suffixes like MW, capacity info
        candidate = re.sub(r'\s+\d+(?:\.\d+)?\s*MW.*$', '', line_stripped).strip()
        candidate = re.sub(r'\s*:.*$', '', candidate).strip()  # Remove colon and after
        
        if _is_known_project(candidate, project_names):
            return candidate
        
        # Try progressive shorter versions
        parts = candidate.split()
        if len(parts) >= 2:
            for i in range(len(parts), 1, -1):
                partial = " ".join(parts[:i])
                if _is_known_project(partial, project_names):
                    return partial
    
    return None


def _is_known_project(candidate: str, project_names: set[str]) -> bool:
    """Check if a candidate string matches a known project name."""
    if not candidate:
        return False
    
    candidate = candidate.strip()
    
    # Exact match
    if candidate in project_names:
        return True
    
    # Case-insensitive match
    candidate_lower = candidate.lower()
    for project_name in project_names:
        if candidate_lower == project_name.lower():
            return True
    
    # Partial match for compound names (e.g., "Taurus A-3" should match if we have project with similar name)
    for project_name in project_names:
        if candidate_lower in project_name.lower() or project_name.lower() in candidate_lower:
            # Additional check to avoid false positives
            if len(candidate) >= 4 and len(project_name) >= 4:  # Minimum meaningful length
                return True
    
    return False


# --- Project mapping (project_name -> project_code) loaded from CSV ---
_PROJECT_NAME_TO_CODE_CACHE: dict[str, str] | None = None


def _default_project_csv_path() -> Path:
    """Resolve the repository-root `data/project.csv` path robustly.

    We locate this file relative to the current file: backend/app/utils.py → repo_root/data/project.csv
    """
    here = Path(__file__).resolve()
    repo_root = here.parents[2]  # .../qenergy-platform
    return repo_root / "data" / "project.csv"


def load_project_name_to_code_mapping(csv_path: Optional[Path] = None, *, force_reload: bool = False) -> dict[str, str]:
    """Load mapping from project_name to project_code from the canonical CSV.

    - CSV schema: project_code;project_name;portfolio_cluster;status
    - Only status==1 considered active and included
    - Keys are normalized to lowercase stripped strings
    - Result is cached in-process; pass force_reload=True to refresh
    """
    global _PROJECT_NAME_TO_CODE_CACHE
    if _PROJECT_NAME_TO_CODE_CACHE is not None and not force_reload:
        return _PROJECT_NAME_TO_CODE_CACHE

    path = csv_path or _default_project_csv_path()
    mapping: dict[str, str] = {}
    if not path.exists():
        _PROJECT_NAME_TO_CODE_CACHE = mapping
        return mapping

    # Try different encodings to handle special characters
    for encoding in ["utf-8", "latin-1", "cp1252"]:
        try:
            with path.open("r", encoding=encoding) as f:
                reader = csv.DictReader(f, delimiter=";")
                for row in reader:
                    try:
                        code = (row.get("project_code") or "").strip()
                        name = (row.get("project_name") or "").strip()
                        status_str = (row.get("status") or "0").strip()
                        if not code or not name:
                            continue
                        if status_str not in ("1", "true", "True"):
                            continue
                        key = name.lower()
                        mapping[key] = code
                    except Exception:
                        # Skip malformed rows silently
                        continue
            break  # Successfully loaded with this encoding
        except UnicodeDecodeError:
            continue  # Try next encoding

    _PROJECT_NAME_TO_CODE_CACHE = mapping
    return mapping


def get_project_code_by_name(project_name: str) -> Optional[str]:
    """Return project_code for a given project_name using CSV mapping.

    Performs case-insensitive lookup after stripping whitespace. Returns None if not found.
    """
    if not project_name:
        return None
    name_key = project_name.strip().lower()
    mapping = load_project_name_to_code_mapping()
    return mapping.get(name_key)


def get_project_code_by_name_db(db, project_name: str) -> Optional[str]:
    """Lookup project_code by project_name in the database (status=1 only)."""
    if not project_name:
        return None
    name_key = project_name.strip().lower()
    rec = db.execute(
        text(
            """
            SELECT project_code
            FROM projects
            WHERE status = 1 AND lower(project_name) = :name
            LIMIT 1
            """
        ),
        {"name": name_key},
    ).first()
    return rec.project_code if rec else None


def seed_projects_from_csv(db, csv_path: Optional[Path] = None, created_by: str = "sys") -> int:
    """Idempotently load projects from CSV into the `projects` table.

    - Inserts or updates rows by `project_code` unique constraint.
    - Returns number of upserts attempted.
    """
    path = csv_path or _default_project_csv_path()
    if not path.exists():
        return 0

    upserts = 0
    # Try different encodings to handle special characters
    for encoding in ["utf-8", "latin-1", "cp1252"]:
        try:
            with path.open("r", encoding=encoding) as f:
                reader = csv.DictReader(f, delimiter=";")
                for row in reader:
                    code = (row.get("project_code") or "").strip()
                    name = (row.get("project_name") or "").strip()
                    portfolio = (row.get("portfolio_cluster") or None)
                    status_raw = (row.get("status") or "0").strip()
                    if not code or not name:
                        continue
                    status_val = 1 if status_raw in ("1", "true", "True") else 0
                    db.execute(
                        text(
                            """
                            INSERT INTO projects (
                                project_code, project_name, portfolio_cluster, status,
                                created_by, updated_by
                            ) VALUES (
                                :project_code, :project_name, :portfolio_cluster, :status,
                                :created_by, :updated_by
                            )
                            ON CONFLICT (project_code) DO UPDATE SET
                                project_name = EXCLUDED.project_name,
                                portfolio_cluster = EXCLUDED.portfolio_cluster,
                                status = EXCLUDED.status,
                                updated_by = EXCLUDED.updated_by,
                                updated_at = NOW()
                            """
                        ),
                        {
                            "project_code": code,
                            "project_name": name,
                            "portfolio_cluster": portfolio,
                            "status": status_val,
                            "created_by": created_by,
                            "updated_by": created_by,
                        },
                    )
                    upserts += 1
            break  # Successfully processed with this encoding
        except UnicodeDecodeError:
            continue  # Try next encoding
    
    db.commit()
    return upserts
